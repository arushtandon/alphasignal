"""Generate AlphaSignal system documentation as a Word (.docx) file."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

OUT = r"c:\Users\tando\Downloads\alphasignal-repo\docs\AlphaSignal-System-Documentation.docx"

doc = Document()

# ── Styles helper ─────────────────────────────────────────────────────────────
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()
    return t

# ── Cover ─────────────────────────────────────────────────────────────────────
add_title('AlphaSignal')
add_p('Quantitative Trading Recommendation System — Technical & Decision-Making Reference', bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f'Document version: v7.9.9  |  Generated: {date.today().strftime("%d %B %Y")}')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('1. Executive Summary')
add_p(
    'AlphaSignal is an autonomous, server-side quantitative stock screening and recommendation '
    'platform. It continuously scans a global equity universe (US, Europe, and Asia index '
    'constituents), computes multi-timeframe technical and fundamental signals for each name, '
    'ranks the highest-conviction opportunities, and publishes up to five Buy and five Sell '
    'recommendations per timeframe on the dashboard. Every recommendation is accompanied by '
    'entry price, Take-Profit (TP1), and Stop-Loss (SL) levels that are horizon-appropriate, '
    'volatility-adjusted, and risk-reward disciplined.'
)
add_p(
    'The system operates independently of whether a user has the dashboard open. A background '
    'scheduler on the Render server runs daily universe scans, regenerates picks on each UTC '
    'trading day, records recommendations into durable trade history, and refreshes open-trade '
    'PnL with live prices and trailing stops.'
)

# ══════════════════════════════════════════════════════════════════════════════
add_h1('2. System Architecture')
add_h2('2.1 Components')
add_bullet('Node.js backend (server.js) — signal engine, API, scheduler, history, backtesting')
add_bullet('Web dashboard (public/index.html) — displays picks, charts, history, exports')
add_bullet('Universe module (universe.js) — index constituent lists for global markets')
add_bullet('Persistent data store (Render disk / DATA_DIR) — picks, history, SL cooldowns')

add_h2('2.2 Data Sources')
add_table(
    ['Source', 'Data Provided', 'Usage'],
    [
        ['Yahoo Finance', 'Daily & weekly OHLCV, live quotes (open/last)', 'All technical indicators, entry pricing'],
        ['Financial Modeling Prep (FMP)', 'S&P 500 / NASDAQ constituents, fundamentals, quality scores', 'Universe, Piotroski, Altman Z, earnings/revenue growth'],
        ['Danelfin API', 'AI Score, Technical, Fundamental, Sentiment, Low Risk subscores', 'Tier overlays, conviction boosts/caps per horizon'],
        ['Alpha Vantage / Finnhub / Yahoo', 'Fallback fundamentals when FMP unavailable', 'International names, P/E, PEG, analyst targets'],
    ]
)

add_h2('2.3 Autonomous Operation Flow')
add_p('The end-to-end pipeline runs as follows:')
steps = [
    '1. Weekly full-universe scan — fetch OHLCV for ~1,000+ tickers across S&P 500, NASDAQ 100, DAX, CAC 40, FTSE 100, Nikkei 225, Nifty 50, HSI',
    '2. Shortlist generation — rank names by composite signal strength; retain top ~120 for deep analysis',
    '3. Daily pick regeneration — on UTC day rollover, re-score shortlist, assign best horizon per ticker, select top 5 Buy + top 5 Sell per pane',
    '4. Live re-pricing — final picks priced at session OPEN (not yesterday\'s close)',
    '5. History recording — new picks saved to durable trade history with frozen entry/levels',
    '6. PnL refresh — open trades updated with trailing SL, signal-flip exits, and $10k notional PnL',
]
for s in steps:
    add_bullet(s)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('3. Timeframes & Strategy Philosophy')
add_table(
    ['Horizon', 'Hold Period', 'Strategy Type', 'Primary Goal'],
    [
        ['Short', '1 day – 1 month (~20 trading days)', 'Mean Reversion', 'Buy statistical discounts at lower SD band / support; fade extensions'],
        ['Medium', '1 – 3 months (~63 trading days)', 'Trend + Pullback', 'Buy pullbacks within established uptrends; short structural breakdowns'],
        ['Long', '4 – 12 months (~180 trading days)', 'SEPA / CANSLIM + Quality', 'Multi-month leaders with MA alignment, earnings growth, fundamental quality'],
    ]
)

add_h2('3.1 Analytical Weighting by Horizon (Display Labels)')
add_table(
    ['Horizon', 'Technical', 'Fundamental', 'News / Sentiment'],
    [
        ['Short', '100%', '—', '—'],
        ['Medium', '70%', '—', '30%'],
        ['Long', '60%', '20%', '20%'],
    ]
)
add_p(
    'Note: These weightings reflect how the UI presents the analysis blend. The actual scoring '
    'engine applies horizon-specific gate logic (see Section 5). Fundamentals are embedded directly '
    'in the long-horizon buy gates (EPS growth, revenue growth, PEG, analyst targets) and via '
    'Danelfin/FMP tier overlays for all horizons where data is available.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('4. Technical Indicators Computed (buildFullTechResult)')
add_p('For every ticker, AlphaSignal computes the following from daily and weekly OHLCV:')

add_h2('4.1 Moving Averages & Trend')
add_table(
    ['Indicator', 'Parameters', 'Role in Decision'],
    [
        ['MA20, MA50, MA100, MA200', 'Simple moving averages', 'Trend regime, support/resistance, SEPA alignment'],
        ['aboveMa20/50/200', 'Price vs MA', 'Bull/bear structure gates; falling-knife detection'],
        ['Golden Cross / Death Cross', 'MA50 vs MA200', 'Medium/long buy and sell regime gates'],
        ['trend20', '20-day price slope', 'Short-term direction; falling-knife filter'],
        ['weeklyTrend', 'Weekly 20-bar slope', 'Medium/long trend confirmation'],
        ['recentTrend', '5-day price change', 'Supplementary momentum context'],
        ['ADX (14)', 'Average Directional Index', 'Trend strength: ranging (<24) vs trending (>28)'],
    ]
)

add_h2('4.2 Momentum & Oscillators')
add_table(
    ['Indicator', 'Parameters', 'Role in Decision'],
    [
        ['RSI (14)', 'Relative Strength Index', 'Tiered thresholds: <20 extreme, <30 oversold, >70 overbought'],
        ['RSI(2)', '2-period RSI', 'Short-horizon mean-reversion timing (washed-out / overbought)'],
        ['Stochastic %K', '14,3', 'Short-horizon oversold/overbought confirmation'],
        ['MACD', '12/26/9 + histogram', 'Trend direction, turning up/down detection'],
        ['Bollinger Bands (%B)', '20-period, 2σ', 'Band position: lower band = discount, upper = extended'],
    ]
)

add_h2('4.3 Volatility & Channels')
add_table(
    ['Indicator', 'Parameters', 'Role in Decision'],
    [
        ['ATR (14)', 'Average True Range', 'Stop sizing, volatility-adaptive floors'],
        ['SD Channels (daily20, daily50, weekly20, weekly50)', 'Linear regression ±1σ/2σ', 'Core timing: buy at lower band, sell at upper band'],
        ['channelPos (buyQuality / sellQuality)', 'Position in channel', 'excellent / good / fair / poor — primary entry quality gate'],
        ['Supertrend', 'ATR-based (10,3 default)', 'Per-horizon trend filter; Strong Buy/Sell on fresh flips'],
        ['supertrendByHz', 'Fast/medium/slow params per horizon', 'Short uses fast ST; long uses slow ST'],
    ]
)

add_h2('4.4 Volume & Structure')
add_table(
    ['Indicator', 'Parameters', 'Role in Decision'],
    [
        ['Volume analysis', '20-day average, relative volume', 'Confirmation: accumulation vs distribution'],
        ['OBV signal', 'On-Balance Volume slope', 'Institutional accumulation gate (medium/long)'],
        ['Swing structure (HH/HL, LH/LL)', '40-bar lookback', 'Bullish/bearish market structure'],
        ['Volume-weighted S/R', '80-bar, 30 clusters', 'Support1/2/3, Resistance1/2/3 with confluence flags'],
        ['Candle patterns', 'Latest bar', 'Bullish/bearish pattern display on charts'],
        ['healthyPullback', 'MA50 + low volume dip', 'Medium-term pullback quality indicator'],
        ['consecutiveLowerCloses', 'Daily count', 'Falling-knife filter for short buys'],
    ]
)

add_h2('4.5 52-Week Range')
add_bullet('high52w / low52w — used in long-horizon CANSLIM "leader near 52-week high" gate')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('5. Fundamental Indicators')
add_p('Fundamentals are fetched from FMP (primary), with Yahoo/Alpha Vantage/Finnhub fallbacks.')

add_h2('5.1 Direct Fundamental Gates (Long Horizon)')
add_table(
    ['Parameter', 'Threshold', 'Buy Gate Impact'],
    [
        ['EPS growth (earningsGrowth)', '≥25%', '+2 gates (CANSLIM)'],
        ['EPS growth', '≥15%', '+1 gate'],
        ['EPS growth', '≥8%', '+0.5 gate'],
        ['Revenue growth', '≥20%', '+1 gate'],
        ['Revenue growth', '≥10%', '+0.5 gate'],
        ['PEG ratio', '0 < PEG < 1.5', '+0.5 gate (GARP)'],
        ['Analyst recommendation', 'Buy / Strong Buy + >15% upside to target', '+1 gate'],
        ['Analyst recommendation', 'Buy + >8% upside', '+0.5 gate'],
        ['Analyst Sell/Strong Sell', '—', 'Buy score × 0.50 penalty'],
        ['EPS decline', '< -10%', '+1 sell gate'],
        ['Revenue decline', '< -8%', '+1 sell gate'],
    ]
)

add_h2('5.2 FMP Quality Score Overlay (applyFmpTierOverlay)')
add_table(
    ['Metric', 'Description', 'Impact'],
    [
        ['FMP Quality Score', 'Composite 0–10 quality rating', 'Tier boost/cap on buy scores per horizon'],
        ['Piotroski F-Score', '0–9 financial strength', 'Short-horizon quality proxy'],
        ['Altman Z-Score', 'Bankruptcy risk', 'Long-horizon quality weighting'],
        ['Analyst Score', 'FMP analyst consensus', 'Medium/long conviction adjustment'],
    ]
)

add_h2('5.3 Danelfin AI Overlay (applyDanelfinTierOverlay)')
add_table(
    ['Danelfin Field', 'Short Horizon Use', 'Medium Horizon Use', 'Long Horizon Use'],
    [
        ['AI Score (aiscore)', 'Tier boost if ≥8 + SD channel', 'Primary tier-1 gate (≥8 + SD + buy track record)', 'Combined with Fundamental for tier-1'],
        ['Technical subscore', 'Boost if ≥7 + buy track record', 'Secondary confirmation', 'Part of composite'],
        ['Fundamental subscore', '—', '—', 'Tier-1 requires Fund ≥7 with AI ≥8'],
        ['Sentiment', 'Display / context', 'News weight proxy', 'News weight proxy'],
        ['Low Risk', '—', '—', '15% weight in long composite (_longQ)'],
        ['buy_track_record', 'Required for tier boosts', 'Required for tier-1', 'Required for tier-1'],
        ['sell_track_record', 'Boosts sell score when AI ≤3', 'Boosts sell on weak AI', 'Boosts sell on weak fundamentals'],
    ]
)

add_h2('5.4 Market Tier Classification')
add_bullet('US large-cap: Full Danelfin + FMP overlays')
add_bullet('European (Danelfin EU): Danelfin + FMP where available')
add_bullet('FMP quality only: FMP overlay without Danelfin')
add_bullet('Technical only: Scores discounted 15% — no reliable fundamental data')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('6. Market Regime Detection')
add_p(
    'Before horizon-specific scoring, each stock is classified into one of three regimes based '
    'on its own technicals (no external index required):'
)
add_table(
    ['Regime', 'Detection Criteria (points ≥4)', 'Buy Multiplier', 'Sell Multiplier'],
    [
        ['Bull', 'Above MA200 + Golden Cross + weekly uptrend + ADX>22 + MACD bull + OBV bull', '×1.30', '×0.50'],
        ['Bear', 'Below MA200 + Death Cross + weekly downtrend + MACD bear + OBV bear + RSI<40', '×0.45 (short) / blocked (long)', '×1.35'],
        ['Neutral', 'Mixed signals', '×1.00 (+1.25 mean-reversion bonus)', '×1.00'],
    ]
)
add_p(
    'Regime directly controls strategy selection: bull markets favour trend-following buys; '
    'bear markets favour breakdown shorts and block long-horizon buys entirely; neutral markets '
    'favour mean-reversion (SD channel timing).'
)

# ══════════════════════════════════════════════════════════════════════════════
add_h1('7. Scoring Engine — computeQuantSignal')
add_p(
    'Each stock receives independent buyScore and sellScore values (0–92) for each of the three '
    'horizons. Scores are built from additive "gates" (weighted evidence points), then mapped to '
    'a 0–92 scale, then adjusted by Supertrend, strict-short rules, and tier overlays.'
)

add_h2('7.1 Short Horizon (Mean Reversion) — Buy Gates')
add_table(
    ['Gate', 'Points', 'Condition'],
    [
        ['SD channel excellent', '+2.5', 'Price at/below lower 1σ band'],
        ['SD channel good', '+1.5', 'Price near lower band'],
        ['At support S1', '+1.3', 'Price within 1.5% of volume-weighted support'],
        ['At MA50 support', '+0.9', 'Within 2% of MA50, above MA50'],
        ['At MA20 support', '+0.6', 'Within 1.5% of MA20'],
        ['RSI(2) < 10', '+1.6', 'Washed out — extreme mean-reversion'],
        ['RSI(2) < 20', '+1.0', 'Oversold fast oscillator'],
        ['Stochastic < 20', '+0.9', 'Oversold confirmation'],
        ['RSI(14) < 30', '+0.9', 'Oversold'],
        ['Lower Bollinger (%B < 10)', '+0.8', 'At lower band'],
        ['MACD turning up', '+0.5', 'Momentum reversal starting'],
        ['OBV bullish', '+0.3', 'Accumulation'],
        ['Ranging market (ADX<24)', '+0.5', 'Mean-reversion favourable'],
        ['Light volume dip (<0.8× avg)', '+0.3', 'Healthy pullback'],
    ]
)

add_h3('Short Buy Penalties (Falling-Knife Filters)')
add_bullet('Strong downtrend (ADX≥28 + downtrend): gates × 0.45')
add_bullet('4+ consecutive lower closes: gates × 0.45')
add_bullet('Below MA20 in downtrend: gates × 0.40')
add_bullet('RSI > 66 or at SD top: capped / penalised')

add_h3('Short Score Mapping')
add_table(
    ['Gates Met', 'Buy Score'],
    [
        ['≥5.5', '90'],
        ['≥4.5', '80'],
        ['≥3.5', '68 (Buy threshold)'],
        ['≥2.8', '58 (Hold)'],
        ['≥2.0', '44'],
        ['<2.0', '≤24'],
    ]
)

add_h2('7.2 Medium Horizon (Trend + Pullback) — Key Buy Gates')
add_table(
    ['Gate', 'Points', 'Condition'],
    [
        ['MA200 bull + Golden Cross', '+2', 'Primary uptrend regime'],
        ['Above MA200', '+1', 'Primary uptrend'],
        ['Golden Cross alone', '+1', 'MA50 > MA200'],
        ['Weekly uptrend', '+1', 'Weekly momentum confirmed'],
        ['ADX ≥28 + above MA200', '+1', 'Strong trending regime'],
        ['SD channel pullback in uptrend', '+1', 'Buy the dip within trend'],
        ['OBV accumulation + RSI 38–68', '+1', 'Institutional demand'],
        ['HH/HL structure + healthy RSI', '+0.7', 'Bullish market structure'],
    ]
)
add_p('Hard disqualifiers: RSI > 74 (×0.50), below MA200 without Golden Cross (×0.20), Death Cross (×0.30), bear regime caps buy at 1.5 gates.')

add_h2('7.3 Long Horizon (SEPA + CANSLIM) — Key Buy Gates')
add_table(
    ['Gate', 'Points', 'Condition'],
    [
        ['SEPA full alignment + rising', '+3', 'Above MA200+MA50, MA50>MA200, weekly uptrend'],
        ['SEPA alignment + Golden Cross', '+2', 'MA stack aligned'],
        ['MA200 + Golden Cross', '+2', 'Structural bull'],
        ['Leader near 52W high', '+1', 'Price ≥75% of 52-week high'],
        ['Weekly uptrend + above MA200', '+1', 'Multi-month momentum'],
        ['ADX ≥25 + above MA200', '+1', 'Trending leader'],
        ['SD channel VCP entry', '+1', 'Volatility contraction pattern at discount'],
        ['RSI 35–65 + above MA200', '+1', 'Healthy momentum zone'],
        ['EPS growth gates', '+0.5 to +2', 'See Section 5.1'],
        ['OBV + HH/HL structure', '+1', 'Accumulation + structure'],
    ]
)
add_p('Hard rule: if price is below MA200, buy score = 0 (no long-term buy in bear structure). Bear regime also sets buy = 0.')

add_h2('7.4 Supertrend Integration (All Horizons)')
add_bullet('Fresh bull flip + buy ≥50: +12 points, labelled "Strong Buy (fresh bull flip)"')
add_bullet('Bull Supertrend + buy ≥55: +6 points confirmation')
add_bullet('Bear Supertrend on medium/long: buy capped at 58 (cannot reach Buy threshold unless fresh flip)')
add_bullet('Fresh bear flip + sell ≥50: +12 points, "Strong Sell"')
add_bullet('Backtest-validated ST (WR ≥60%): additional +2 confidence boost')

add_h2('7.5 Strict Selective Short Rules')
add_p('A sell score ≥62 is demoted to ≤55 (Hold) unless ALL of the following are true:')
add_bullet('Price below MA200 (structural downtrend)')
add_bullet('Regime is not bull')
add_bullet('Horizon Supertrend is bearish')
add_bullet('Real distribution: OBV falling, bearish structure, RSI falling, or bearish volume')
add_bullet('Fundamentals not strong (EPS <15% AND revenue <12% growth)')
add_bullet('Not within 5 days of earnings (squeeze risk)')

add_h2('7.6 Action & Rating Thresholds (deriveActionRating)')
add_table(
    ['Condition', 'Action', 'Rating'],
    [
        ['buyScore ≥ 78 (and buy ≥ sell)', 'Buy', 'Strong Buy'],
        ['buyScore ≥ 62 (and buy ≥ sell)', 'Buy', 'Buy'],
        ['buyScore < 62', 'Hold', 'Hold'],
        ['sellScore ≥ 74 (and sell > buy)', 'Sell', 'Strong Sell'],
        ['sellScore ≥ 62 (and sell > buy)', 'Sell', 'Sell'],
    ]
)

add_h2('7.7 Structural Overrides (applyTierScoreCaps)')
add_bullet('Falling knife (below MA20 & MA200, RSI<45): buy capped at 45 — blocked')
add_bullet('Below MA200 short-term buy: capped at 61 (cannot reach Buy)')
add_bullet('Below MA20 & MA50 short-term: capped at 61')
add_bullet('Medium below MA50 in downtrend: capped at 73 (no Strong Buy)')
add_bullet('Tier 0 stocks: buy capped at 72; Tier 1 capped at 88')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('8. Recommendation Selection (Dashboard Picks)')
add_p('After scoring the shortlist (~120 names), the server selects final dashboard picks:')

add_h2('8.1 Cross-Timeframe Deduplication ("Best Conviction Wins")')
add_p(
    'Each ticker is assigned to exactly ONE horizon — the one where it has the highest '
    'qualifying score. A name cannot appear in Short, Medium, and Long simultaneously.'
)

add_h2('8.2 Selection Criteria per Pane')
add_bullet('Action must be Buy (for buy panes) or Sell (for sell panes)')
add_bullet('Score must be ≥ 62 (buyScore or sellScore depending on direction)')
add_bullet('Valid entry and stop-loss levels must exist (hasPx check)')
add_bullet('Not in SL cooldown (6-day block after stop-out on same ticker+horizon)')
add_bullet('No-repeat filter: if already open in same direction, suppress re-recommendation')

add_h2('8.3 Top-N Selection')
add_p('Top 5 names per pane by score: short buys, medium buys, long buys, short sells, medium sells, long sells.')

add_h2('8.4 Entry Price')
add_p(
    'Entry = session OPEN price for the recommendation day (what a trader would get filled at '
    'on the open). Falls back to last live price only if open is unavailable. Entry is frozen '
    'once recorded in history and never overwritten.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('9. Take-Profit & Stop-Loss Logic')
add_p(
    'TP/SL levels are computed server-side and enforced consistently across the dashboard, '
    'trade history, PnL refresh, and backtesting. Levels scale with the holding period.'
)

add_h2('9.1 Horizon Minimum Floors (HORIZON_MIN_PCT — v7.9.9)')
add_table(
    ['Horizon', 'Minimum SL Distance', 'Minimum TP1 Distance', 'Minimum R:R'],
    [
        ['Short', '3%', '6%', '1.8×'],
        ['Medium', '8%', '15%', '1.85×'],
        ['Long', '16%', '30%', '1.85×'],
    ]
)
add_p(
    'These floors ensure a long-term trade never shows a 2–5% noise stop. Volatile stocks '
    'receive wider stops via ATR multiples (short 2.0×, medium 4.0×, long 6.5× ATR) on top '
    'of the percentage floor.'
)

add_h2('9.2 Short Horizon — Mean Reversion Levels')
add_p('Strategy: defined target at SD channel mean / nearest resistance; stop beyond opposite band.')
add_bullet('Target: minimum of (channel mean, resistance1), floored to horizon min-TP1')
add_bullet('Stop: 1.5× ATR below entry or channel lower-2σ, with min 1.2× ATR distance')
add_bullet('Exit: full position closed at target or stop (no trailing)')

add_h2('9.3 Medium & Long — Hybrid Exit (Trailing + TP1 Partial)')
add_p('Strategy: book partial profit at TP1, trail the remainder with a dynamic stop.')
add_bullet('TP1 = minRR × stop distance (reward ≥ 1.85× risk), floored to horizon min-TP1')
add_bullet('At TP1 hit: 50% of position booked at TP1 price; remainder stop moved to breakeven')
add_bullet('Post-TP1: daily % ratchet trail — stop moves up on favourable days, never loosens')
add_bullet('Pre-TP1 stop: trailing stop from technicals (channel, S/R, MA, Supertrend), ratcheted daily EOD')
add_bullet('No fixed TP2 — remainder rides until trailing stop hit or time limit')

add_h2('9.4 Trailing Stop Calculation (computeTrailingStopFromTech)')
add_p('Stop candidates drawn from (long example):')
add_bullet('Weekly channel lower-2σ, weekly-50 lower-1σ, MA200 × 0.96')
add_bullet('Supertrend value (if bullish and below entry)')
add_bullet('Final stop = widest valid candidate, floored to minGap (horizon % or ATR multiple)')

add_h2('9.5 Risk-Reward Enforcement (applyHorizonMinPctFloors)')
add_p(
    'After raw levels are computed, enforceMinRiskReward widens TP1 if reward < minRR × risk. '
    'This guarantees positive expectancy: every trade risks X% to make at least 1.85X%.'
)

add_h2('9.6 Exit Conditions Summary')
add_table(
    ['Exit Type', 'When Triggered', 'Horizons'],
    [
        ['TP1 Hit (partial)', 'Price reaches TP1', 'Medium, Long (50% booked)'],
        ['SL Hit', 'Price breaches trailing stop', 'All'],
        ['TP1 then SL', 'Post-TP1 remainder stopped at ratcheted level', 'Medium, Long'],
        ['Signal Exit (reversal)', 'Live signal flips Buy→Sell or Sell→Buy (score ≥62 opposite)', 'All (reversal-only policy)'],
        ['Time Limit', 'Full hold period elapsed without exit', 'All (20 / 63 / 180 days)'],
        ['Mean-reversion target', 'Price reaches channel mean/resistance', 'Short only'],
    ]
)

add_h2('9.7 SL Cooldown')
add_p(
    'After a stop-loss is hit, the same ticker+horizon is blocked from new Buy recommendations '
    'for 6 days. This prevents re-entering a name that just failed.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('10. Trade History & PnL')
add_h2('10.1 Recording')
add_bullet('Every daily pick is automatically saved to server-side durable history')
add_bullet('Entry price, date, TP/SL levels frozen on first record — never overwritten')
add_bullet('Same-day re-scans do not duplicate or re-price existing trades')

add_h2('10.2 PnL Calculation')
add_bullet('Fixed $10,000 notional per trade for dollar PnL display')
add_bullet('Unrealized PnL: mark-to-market for open trades using live quotes')
add_bullet('Realized PnL: computed on closed trades (TP, SL, signal exit, time limit)')
add_bullet('Buy-side vs Sell-side PnL breakdown shown separately in history summary')

add_h2('10.3 Win Rate Definition')
add_p(
    'A closed trade counts as a WIN unless it hit stop-loss (sl_hit) or has negative PnL. '
    'TP1 partial hits, signal exits with profit, and time-limit exits with profit all count as wins.'
)

add_h2('10.4 Signal-Flip Exit Policy (Reversal-Only)')
add_p(
    'Open trades are automatically closed only on a true directional reversal: a Buy trade closes '
    'when the live sellScore reaches ≥62 (Sell signal), and vice versa. Softening to Hold does NOT '
    'trigger an exit. Same-day flips are blocked (trade must be at least one day old).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('11. Backtesting')
add_p(
    'Backtesting uses the same signal engine (computeQuantSignal) and exit simulators '
    '(simulateHybridExit / simulateMeanReversionExit) as live trading, ensuring reported '
    'win rates match what the live rules would produce.'
)
add_bullet('Data: up to 5 years of daily OHLCV with bounded 320-bar technical windows')
add_bullet('Walk-forward: signals computed bar-by-bar without look-ahead bias')
add_bullet('Exit: identical hybrid/mean-reversion logic with trailing stops and TP1 partials')
add_bullet('Supertrend backtest WR stored per ticker for confidence overlay')
add_bullet('Results: win rate, average return, trade count per horizon displayed on dashboard cards')

# ══════════════════════════════════════════════════════════════════════════════
add_h1('12. Complete Decision Flow (End-to-End)')
flow = [
    ('Step 1 — Data Ingestion', 'Fetch 2 years daily + weekly OHLCV from Yahoo Finance for each universe ticker.'),
    ('Step 2 — Technical Analysis', 'Compute 30+ indicators: MAs, RSI, MACD, ADX, ATR, SD channels, Supertrend, S/R, OBV, Bollinger, Stochastic, candle patterns, swing structure.'),
    ('Step 3 — Regime Classification', 'Score bull/bear/neutral points from MA alignment, crosses, weekly trend, ADX, MACD, OBV, RSI.'),
    ('Step 4 — Horizon Scoring', 'Apply horizon-specific gate logic (mean-reversion for short, trend+pullback for medium, SEPA+CANSLIM for long). Accumulate buyGates/sellGates → map to buyScore/sellScore.'),
    ('Step 5 — Supertrend Filter', 'Boost/cap scores based on per-horizon Supertrend direction and fresh flips.'),
    ('Step 6 — Strict Short Gate', 'Demote sell scores that lack bearish confluence (below MA200 + ST bear + distribution).'),
    ('Step 7 — Fundamental Fetch', 'Retrieve FMP/Danelfin/Yahoo fundamentals: EPS growth, revenue, PEG, Piotroski, Altman Z, analyst targets.'),
    ('Step 8 — Tier Overlays', 'Apply Danelfin AI Score and FMP quality overlays. Upgrade tier (0→1→2) for high-conviction names. Cap scores for weak fundamentals.'),
    ('Step 9 — Structural Overrides', 'Block falling knives, cap counter-trend buys, enforce MA-based hard limits.'),
    ('Step 10 — Action Derivation', 'Map final scores to Buy/Hold/Sell and Strong Buy/Strong Sell using thresholds (62/78 buy, 62/74 sell).'),
    ('Step 11 — SL Cooldown Check', 'Suppress recommendations for tickers recently stopped out.'),
    ('Step 12 — Price Level Calculation', 'Compute entry (session open), TP1, and SL using horizon-appropriate channel/S/R/ATR math with min-% floors and R:R enforcement.'),
    ('Step 13 — Cross-Timeframe Dedup', 'Assign each ticker to its single best horizon by highest qualifying score.'),
    ('Step 14 — Top-5 Selection', 'Select top 5 Buy and top 5 Sell per pane. Apply no-repeat filter for already-open positions.'),
    ('Step 15 — Dashboard Publish', 'Serve picks to dashboard. Record in history. Begin PnL tracking with trailing stops and reversal-only exit monitoring.'),
]
for title, desc in flow:
    add_h3(title)
    add_p(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
add_h1('13. Glossary')
glossary = [
    ('SD Channel', 'Standard Deviation regression channel — price relative to ±1σ/2σ bands from a linear regression mean'),
    ('Gate', 'A weighted evidence point in the scoring engine; more gates = higher conviction'),
    ('Tier', 'Quality classification: Tier 0 = standard, Tier 1 = Danelfin≥8 + SD channel, Tier 2 = + catalyst'),
    ('VCP', 'Volatility Contraction Pattern — price consolidating near channel mean before breakout'),
    ('SEPA', 'Specific Entry Point Analysis — Mark Minervini-style MA alignment for leading stocks'),
    ('CANSLIM', 'William O\'Neil methodology — earnings growth, revenue growth, leader near highs'),
    ('Hybrid Exit', 'Book 50% at TP1, trail remainder with ratcheting stop — captures partial win + lets winners run'),
    ('Mean Reversion', 'Strategy of buying statistical discounts and selling extensions in ranging markets'),
    ('Supertrend', 'ATR-based trend indicator — price above = bull, below = bear; flips signal reversals'),
    ('R:R (Risk-Reward)', 'Ratio of potential profit (TP1 distance) to potential loss (SL distance)'),
    ('SL Cooldown', '6-day block on re-recommending a ticker after its stop-loss was hit'),
    ('Falling Knife', 'Stock in active decline below key MAs with weak RSI — buy blocked'),
    ('PnL Notional', 'Fixed $10,000 position size used for dollar PnL display in history'),
]
add_table(['Term', 'Definition'], glossary)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
add_p(
    'This document describes the AlphaSignal system as implemented in server.js and '
    'public/index.html (build v7.9.9). Thresholds, floors, and overlay weights may be '
    'tuned in future releases; the architecture and decision flow remain consistent.',
    bold=False
)

doc.save(OUT)
print(f'Created: {OUT}')
